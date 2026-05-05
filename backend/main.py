"""
DocScanPro Backend — FastAPI v3
Capabilities: pytesseract OCR, Mistral AI OCR (tables/charts),
              pdf2image, OCRmyPDF, python-docx, reportlab PDF
"""

from fastapi import FastAPI, File, UploadFile, Form, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pathlib import Path
from typing import List
import io, os, re, uuid, tempfile, datetime, base64, httpx

# ── Optional library imports ────────────────────────────────────────────────
try:
    import pytesseract
    from PIL import Image
    # Free memory after import
    import gc
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import ocrmypdf
    OCRMYPDF_AVAILABLE = True
except ImportError:
    OCRMYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        HRFlowable, Table, TableStyle,
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# ── App ─────────────────────────────────────────────────────────────────────
app = FastAPI(title="DocScanPro API", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=True,
)

TEMP_DIR = Path(tempfile.gettempdir()) / "docscanner"
TEMP_DIR.mkdir(exist_ok=True)
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif", ".gif"}
MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "")
GOOGLE_VISION_CREDENTIALS = os.environ.get("GOOGLE_VISION_CREDENTIALS", "")

# ── Google Vision setup ───────────────────────────────────────────────────────
VISION_CLIENT = None
VISION_AVAILABLE = False

if GOOGLE_VISION_CREDENTIALS:
    try:
        import json as _json
        from google.cloud import vision
        from google.oauth2 import service_account
        _creds_dict = _json.loads(GOOGLE_VISION_CREDENTIALS)
        _creds = service_account.Credentials.from_service_account_info(
            _creds_dict,
            scopes=["https://www.googleapis.com/auth/cloud-platform"]
        )
        VISION_CLIENT = vision.ImageAnnotatorClient(credentials=_creds)
        VISION_AVAILABLE = True
    except Exception as e:
        print(f"Google Vision setup failed: {e}")

# ── Google Vision usage tracking (in-memory, resets on server restart) ────────
# For persistent tracking, this should be in Supabase — works fine for now
import datetime as _dt
_vision_usage = {"month": _dt.date.today().strftime("%Y-%m"), "count": 0}
VISION_MONTHLY_LIMIT = 1000  # free tier limit

def vision_quota_ok() -> bool:
    """Check if we're within Google Vision free quota."""
    today_month = _dt.date.today().strftime("%Y-%m")
    if _vision_usage["month"] != today_month:
        _vision_usage["month"] = today_month
        _vision_usage["count"] = 0
    return _vision_usage["count"] < VISION_MONTHLY_LIMIT

def vision_increment():
    _vision_usage["count"] += 1


# ── Google Vision OCR ─────────────────────────────────────────────────────────
def google_vision_ocr(image_bytes: bytes) -> str:
    """Fast, accurate OCR using Google Cloud Vision."""
    if not VISION_AVAILABLE or not VISION_CLIENT:
        raise ValueError("Google Vision not available")
    if not vision_quota_ok():
        raise ValueError(f"Google Vision quota reached ({VISION_MONTHLY_LIMIT}/month) — using fallback")

    from google.cloud import vision
    image   = vision.Image(content=image_bytes)
    response= VISION_CLIENT.document_text_detection(image=image)

    if response.error.message:
        raise ValueError(f"Vision API error: {response.error.message}")

    vision_increment()

    # Extract full text with layout preservation
    full_text = response.full_text_annotation.text if response.full_text_annotation else ""
    return full_text.strip()

# ── Helpers ──────────────────────────────────────────────────────────────────
def tmp(suffix="") -> Path:
    return TEMP_DIR / f"{uuid.uuid4().hex}{suffix}"

def rm(path):
    try: os.unlink(path)
    except: pass

def safe_name(s: str, max_len=60) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", s)[:max_len] or "document"


# ── Mistral AI OCR ────────────────────────────────────────────────────────────
async def mistral_ocr(image_bytes: bytes, mime_type: str = "image/jpeg") -> str:
    if not MISTRAL_API_KEY:
        raise ValueError("MISTRAL_API_KEY not set")

    b64 = base64.b64encode(image_bytes).decode()
    data_url = f"data:{mime_type};base64,{b64}"

    payload = {
        "model": "pixtral-12b-2409",
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": data_url}},
                {"type": "text", "text": (
                    "You are a precise OCR engine for audit document processing. "
                    "Extract ALL text from this document with perfect accuracy.\n\n"
                    "RULES:\n"
                    "1. TABLES: render as markdown tables using | separators with header rows\n"
                    "2. CHARTS/GRAPHS: describe as [Chart: type, title, key values]\n"
                    "3. FORMS: render as 'Field: Value' pairs\n"
                    "4. Preserve ALL numbers, dates, amounts exactly\n"
                    "5. Preserve headings, numbered lists, bullets\n"
                    "6. Hindi/Bengali/Assamese: extract in original script\n"
                    "7. Return ONLY extracted content — no commentary, no backticks"
                )}
            ]
        }],
        "max_tokens": 4096,
        "temperature": 0,
    }

    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(
            "https://api.mistral.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"},
            json=payload,
        )
        if not r.is_success:
            raise ValueError(f"Mistral API error {r.status_code}: {r.text[:200]}")
        return r.json()["choices"][0]["message"]["content"].strip()


# ── Structure detection ───────────────────────────────────────────────────────
def detect_structure(text: str) -> list[dict]:
    lines = text.split("\n")
    sections: list[dict] = []
    cur = {"title": "", "level": 0, "type": "preamble", "content": []}

    for line in lines:
        s = line.strip()
        if not s:
            cur["content"].append("")
            continue

        level, typ = None, None
        if re.match(r"^(CHAPTER|PART|TITLE|APPENDIX)\s+[\dIVXivxa-z]", s, re.I):
            level, typ = 1, "chapter"
        elif re.match(r"^#{1,2}\s+\S", s):
            level, typ = 1, "chapter"
        elif s.isupper() and 7 <= len(s) <= 90 and sum(c.isalpha() for c in s) > len(s)*0.55:
            level, typ = 1, "chapter"
        elif re.match(r"^#{3,4}\s+\S", s):
            level, typ = 2, "section"
        elif re.match(r"^(SECTION|ARTICLE|CLAUSE)\s+\d", s, re.I):
            level, typ = 2, "section"
        elif re.match(r"^\d{1,2}[.)]\s+[A-Z]", s):
            level, typ = 2, "section"
        elif re.match(r"^\d{1,2}\.\d{1,3}[.):\s]", s):
            level, typ = 3, "subsection"

        if level is not None:
            if cur["title"] or cur["content"]:
                sections.append(cur)
            cur = {"title": re.sub(r"^#+\s+", "", s), "level": level, "type": typ, "content": []}
        else:
            cur["content"].append(s)

    if cur["title"] or cur["content"]:
        sections.append(cur)
    return sections


# ── Markdown table parser ─────────────────────────────────────────────────────
def parse_markdown_table(lines: list[str]) -> list[list[str]] | None:
    rows = []
    for line in lines:
        if re.match(r"^\|[\s\-:|]+\|", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows if len(rows) >= 1 else None


# ── PDF builder ───────────────────────────────────────────────────────────────
def build_pdf(sections: list[dict], title: str, author: str, out: str):
    TEAL    = colors.HexColor("#0F6E56")
    TEAL_L  = colors.HexColor("#D1FAE5")
    TEAL_XL = colors.HexColor("#F0FDF9")
    DARK    = colors.HexColor("#111827")
    GRAY    = colors.HexColor("#6B7280")
    WHITE   = colors.white

    pw, ph = A4
    mg = 2.5 * cm

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(TEAL)
        canvas.rect(0, ph - 1.2*cm, pw, 1.2*cm, fill=1, stroke=0)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(WHITE)
        canvas.drawString(mg, ph - 0.8*cm, (title or "")[:60])
        canvas.drawRightString(pw - mg, ph - 0.8*cm,
                               datetime.date.today().strftime("%d %b %Y"))
        canvas.setFillColor(TEAL_L)
        canvas.rect(0, 0, pw, 1.0*cm, fill=1, stroke=0)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(TEAL)
        canvas.drawString(mg, 0.35*cm, author or "DocScanPro · Audit Edition")
        canvas.drawRightString(pw - mg, 0.35*cm, f"Page {doc.page}")
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        out, pagesize=A4,
        leftMargin=mg, rightMargin=mg,
        topMargin=mg + 0.5*cm, bottomMargin=mg + 0.2*cm,
        title=title, author=author or "DocScanPro",
    )

    # Styles
    cover_h   = ParagraphStyle("cover_h",   fontSize=28, fontName="Helvetica-Bold",
                                textColor=TEAL, alignment=TA_CENTER, spaceAfter=10)
    cover_sub = ParagraphStyle("cover_sub", fontSize=11, fontName="Helvetica",
                                textColor=GRAY, alignment=TA_CENTER, spaceAfter=4)
    cover_tag = ParagraphStyle("cover_tag", fontSize=9, fontName="Helvetica-Bold",
                                textColor=WHITE, backColor=TEAL, alignment=TA_CENTER,
                                borderPad=8, spaceAfter=4)
    ch_h      = ParagraphStyle("ch_h",  fontSize=18, fontName="Helvetica-Bold",
                                textColor=TEAL, spaceBefore=6, spaceAfter=8)
    sec_h     = ParagraphStyle("sec_h", fontSize=13, fontName="Helvetica-Bold",
                                textColor=DARK, spaceBefore=12, spaceAfter=5)
    sub_h     = ParagraphStyle("sub_h", fontSize=11, fontName="Helvetica-Bold",
                                textColor=colors.HexColor("#374151"),
                                spaceBefore=8, spaceAfter=3)
    body      = ParagraphStyle("body",  fontSize=9.5, fontName="Helvetica",
                                textColor=DARK, leading=15, spaceAfter=4,
                                alignment=TA_JUSTIFY)
    bullet_s  = ParagraphStyle("bullet", fontSize=9.5, fontName="Helvetica",
                                textColor=DARK, leading=14, spaceAfter=2,
                                leftIndent=14)
    chart_cap = ParagraphStyle("chart_cap", fontSize=9, fontName="Helvetica-Oblique",
                                textColor=GRAY, alignment=TA_CENTER, spaceAfter=8)
    toc_ch    = ParagraphStyle("toc_ch", fontSize=11, fontName="Helvetica-Bold",
                                textColor=DARK, spaceBefore=5, spaceAfter=1)
    toc_sec   = ParagraphStyle("toc_sec", fontSize=10, fontName="Helvetica",
                                textColor=GRAY, leftIndent=14, spaceAfter=1)
    toc_sub   = ParagraphStyle("toc_sub", fontSize=9, fontName="Helvetica",
                                textColor=GRAY, leftIndent=28, spaceAfter=1)

    story = []

    # Cover
    story.append(Spacer(1, 3.5*cm))
    story.append(HRFlowable(width="100%", thickness=4, color=TEAL, spaceAfter=20))
    story.append(Paragraph(title or "Untitled Document", cover_h))
    story.append(HRFlowable(width="100%", thickness=1, color=TEAL_L, spaceAfter=18))
    if author:
        story.append(Paragraph(f"Prepared by: {author}", cover_sub))
    story.append(Paragraph(
        datetime.datetime.now().strftime("Generated: %B %d, %Y at %H:%M"), cover_sub))
    story.append(Spacer(1, 0.6*cm))
    story.append(Paragraph("OCR PROCESSED — AUDIT DOCUMENT", cover_tag))
    story.append(PageBreak())

    # TOC
    toc_entries = [(s["title"], s["level"]) for s in sections if s.get("title")]
    if toc_entries:
        story.append(Paragraph("Table of Contents", ch_h))
        story.append(HRFlowable(width="100%", thickness=0.5, color=TEAL_L, spaceAfter=10))
        ch_n = sec_n = sub_n = 0
        for ttl, lvl in toc_entries:
            if lvl == 1:
                ch_n += 1; sec_n = 0; sub_n = 0
                story.append(Paragraph(f"{ch_n}.  {ttl}", toc_ch))
            elif lvl == 2:
                sec_n += 1; sub_n = 0
                story.append(Paragraph(f"{ch_n}.{sec_n}  {ttl}", toc_sec))
            elif lvl == 3:
                sub_n += 1
                story.append(Paragraph(f"{ch_n}.{sec_n}.{sub_n}  {ttl}", toc_sub))
        story.append(PageBreak())

    def make_table(rows):
        max_cols = max(len(r) for r in rows)
        norm = [r + [""] * (max_cols - len(r)) for r in rows]
        cs  = ParagraphStyle("cs",  fontSize=8.5, fontName="Helvetica", leading=12, textColor=DARK)
        hs  = ParagraphStyle("hs",  fontSize=8.5, fontName="Helvetica-Bold", leading=12, textColor=WHITE)
        td  = [[Paragraph(str(c).replace("&","&amp;").replace("<","&lt;"),
                          hs if ri==0 else cs) for c in r] for ri, r in enumerate(norm)]
        cw  = (pw - mg*2) / max_cols
        t   = Table(td, colWidths=[cw]*max_cols, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,0),  TEAL),
            ("FONTNAME",     (0,0),(-1,0),  "Helvetica-Bold"),
            ("BOTTOMPADDING",(0,0),(-1,0),  7),
            ("TOPPADDING",   (0,0),(-1,0),  7),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE, TEAL_XL]),
            ("FONTSIZE",     (0,1),(-1,-1),  8.5),
            ("TOPPADDING",   (0,1),(-1,-1),  5),
            ("BOTTOMPADDING",(0,1),(-1,-1),  5),
            ("LEFTPADDING",  (0,0),(-1,-1),  6),
            ("RIGHTPADDING", (0,0),(-1,-1),  6),
            ("GRID",         (0,0),(-1,-1),  0.4, colors.HexColor("#D1D5DB")),
            ("BOX",          (0,0),(-1,-1),  0.8, TEAL),
            ("VALIGN",       (0,0),(-1,-1),  "MIDDLE"),
        ]))
        return t

    def render_block(lines):
        i = 0
        while i < len(lines):
            line = lines[i]
            if not line.strip():
                story.append(Spacer(1, 3)); i += 1; continue

            if line.strip().startswith("|"):
                tb = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    tb.append(lines[i]); i += 1
                rows = parse_markdown_table(tb)
                if rows:
                    story.append(make_table(rows))
                    story.append(Spacer(1, 6))
                continue

            if re.match(r"^\[?(chart|graph|figure|diagram)\b", line, re.I):
                story.append(Paragraph(f"📊 {line}", chart_cap)); i += 1; continue

            if re.match(r"^[\-\*•]\s+", line):
                txt = re.sub(r"^[\-\*•]\s+", "", line)
                story.append(Paragraph(f"• {txt}", bullet_s)); i += 1; continue

            if re.match(r"^\d+[.)]\s+", line):
                story.append(Paragraph(line, bullet_s)); i += 1; continue

            buf = []
            while i < len(lines) and lines[i].strip() and \
                  not lines[i].strip().startswith("|") and \
                  not re.match(r"^[\-\*•\d]", lines[i]):
                buf.append(lines[i].strip()); i += 1
            if buf:
                txt = " ".join(buf).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                story.append(Paragraph(txt, body))

    for i, sec in enumerate(sections):
        lvl, title_, lines = sec.get("level",0), sec.get("title",""), sec.get("content",[])
        if title_:
            if lvl <= 1:
                if i > 0: story.append(PageBreak())
                story.append(HRFlowable(width="100%", thickness=2, color=TEAL, spaceAfter=4))
                story.append(Paragraph(title_, ch_h))
                story.append(HRFlowable(width="100%", thickness=0.5, color=TEAL_L, spaceAfter=6))
            elif lvl == 2:
                story.append(Paragraph(title_, sec_h))
            else:
                story.append(Paragraph(title_, sub_h))
        render_block(lines)

    pdf_doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


# ── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx(sections: list[dict], title: str, author: str, out: str):
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3.0); sec.right_margin=Cm(2.5)

    TEAL_RGB  = RGBColor(0x0F, 0x6E, 0x56)
    GRAY_RGB  = RGBColor(0x6B, 0x72, 0x80)
    WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)

    h = doc.add_heading(title or "Untitled Document", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs: h.runs[0].font.color.rgb = TEAL_RGB

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(datetime.datetime.now().strftime("Generated: %B %d, %Y"))
    r.font.color.rgb = GRAY_RGB; r.font.size = Pt(10)
    if author:
        pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(f"Prepared by: {author}")
        ra.font.color.rgb = GRAY_RGB; ra.font.size = Pt(10)

    doc.add_page_break()
    tp = doc.add_paragraph("Table of Contents"); tp.style = "Heading 1"
    if tp.runs: tp.runs[0].font.color.rgb = TEAL_RGB
    doc.add_paragraph("[Update TOC: References → Update Table]").runs[0].italic = True
    doc.add_page_break()

    def add_table(rows):
        if not rows: return
        mc = max(len(r) for r in rows)
        norm = [r + [""]*( mc-len(r)) for r in rows]
        t = doc.add_table(rows=len(norm), cols=mc); t.style = "Table Grid"
        for ri, row in enumerate(norm):
            for ci, ct in enumerate(row):
                cell = t.cell(ri, ci); cell.text = str(ct)
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs \
                      else cell.paragraphs[0].add_run(str(ct))
                if ri == 0:
                    run.font.bold = True; run.font.color.rgb = WHITE_RGB
                    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
                    shd.set(qn("w:fill"),"0F6E56"); tcPr.append(shd)
        doc.add_paragraph()

    for i, sec in enumerate(sections):
        lvl, title_, lines = sec.get("level",0), sec.get("title",""), sec.get("content",[])
        if title_:
            hh = doc.add_heading(title_, level=min(max(lvl,1),4))
            if hh.runs and lvl<=1: hh.runs[0].font.color.rgb = TEAL_RGB
            if lvl<=1 and i>0: doc.add_page_break()
        j = 0
        while j < len(lines):
            line = lines[j]
            if line.strip().startswith("|"):
                tb = []
                while j < len(lines) and lines[j].strip().startswith("|"):
                    tb.append(lines[j]); j += 1
                rows = parse_markdown_table(tb)
                if rows: add_table(rows)
            elif line.strip():
                doc.add_paragraph(line.strip()); j += 1
            else:
                j += 1
    doc.save(out)


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/api/health")
@app.head("/api/health")
def health():
    return {
        "status": "ok",
        "capabilities": {
            "ocr_pytesseract":   OCR_AVAILABLE,
            "ocr_mistral":       bool(MISTRAL_API_KEY),
            "ocr_google_vision": VISION_AVAILABLE,
            "vision_quota_ok":   vision_quota_ok(),
            "vision_used":       _vision_usage["count"],
            "vision_limit":      VISION_MONTHLY_LIMIT,
            "pdf2image":         PDF2IMAGE_AVAILABLE,
            "ocrmypdf":          OCRMYPDF_AVAILABLE,
            "docx_export":       DOCX_AVAILABLE,
            "pdf_export":        REPORTLAB_AVAILABLE,
        },
    }


@app.get("/api/vision/usage")
def vision_usage():
    """Check Google Vision quota usage."""
    return {
        "month":    _vision_usage["month"],
        "used":     _vision_usage["count"],
        "limit":    VISION_MONTHLY_LIMIT,
        "remaining":VISION_MONTHLY_LIMIT - _vision_usage["count"],
        "enabled":  VISION_AVAILABLE and vision_quota_ok(),
        "pct":      round(_vision_usage["count"] / VISION_MONTHLY_LIMIT * 100, 1),
    }


@app.post("/api/ocr")
async def ocr(
    file:     UploadFile = File(...),
    language: str        = Form("eng"),
    engine:   str        = Form("auto"),
):
    data = await file.read()
    ext  = Path(file.filename or "").suffix.lower()

    images_bytes: list[bytes] = []
    if ext == ".pdf":
        if not PDF2IMAGE_AVAILABLE:
            raise HTTPException(400, "pdf2image not available")
        for img in convert_from_bytes(data, dpi=150, fmt="jpeg"):
            buf = io.BytesIO()
            img.save(buf, "JPEG", quality=85)
            images_bytes.append(buf.getvalue())
            del img
    elif ext in IMAGE_EXTS or (file.content_type or "").startswith("image/"):
        images_bytes = [data]
    else:
        raise HTTPException(400, f"Unsupported: {ext}")

    page_texts  = []
    engine_used = "tesseract"

    for i, img_bytes in enumerate(images_bytes):
        text = ""

        # ── Engine selection priority: Vision → Mistral → Tesseract ──────────
        if engine in ("auto", "vision") and VISION_AVAILABLE and vision_quota_ok():
            try:
                text        = google_vision_ocr(img_bytes)
                engine_used = "google_vision"
            except Exception as e:
                print(f"Google Vision failed (page {i+1}): {e} — falling back")

        if not text and (engine in ("auto", "mistral") and bool(MISTRAL_API_KEY)):
            try:
                text        = await mistral_ocr(img_bytes)
                engine_used = "mistral"
            except Exception as e:
                print(f"Mistral failed (page {i+1}): {e} — falling back")

        if not text:
            if not OCR_AVAILABLE:
                raise HTTPException(500, "No OCR engine available")
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode not in ("RGB","L","RGBA"): img = img.convert("RGB")
            text        = pytesseract.image_to_string(img, lang=language)
            engine_used = "tesseract"
            del img

        import gc; gc.collect()
        page_texts.append({"page": i+1, "text": text.strip()})

    full_text = "\n\n--- Page Break ---\n\n".join(p["text"] for p in page_texts)
    sections  = detect_structure(full_text)

    return JSONResponse({
        "text": full_text, "pages": len(images_bytes),
        "page_texts": page_texts, "sections": sections,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "engine_used": engine_used,
        "vision_remaining": VISION_MONTHLY_LIMIT - _vision_usage["count"],
    })


@app.post("/api/export/searchable-pdf")
async def export_searchable_pdf(
    background_tasks: BackgroundTasks,
    language: str = Form("eng"),
    files: List[UploadFile] = File(...),
):
    """
    Takes one or more images, returns a searchable PDF with text layer.
    Uses ocrmypdf (which uses pikepdf internally — no extra deps needed).
    """
    if not OCRMYPDF_AVAILABLE:
        raise HTTPException(500, "ocrmypdf not installed")

    page_paths = []
    for f in files:
        data = await f.read()
        ext  = Path(f.filename or "img.jpg").suffix.lower() or ".jpg"
        p    = tmp(ext)
        p.write_bytes(data)
        page_paths.append(p)

    out_path = tmp(".pdf")

    try:
        if len(page_paths) == 1:
            # Single page
            ocrmypdf.ocr(
                str(page_paths[0]), str(out_path),
                language=language, deskew=True, clean=True,
                optimize=1, force_ocr=True, progress_bar=False,
            )
        else:
            # Multi-page: OCR each page separately then merge with pikepdf
            import pikepdf
            page_pdfs = []
            for pp in page_paths:
                page_out = tmp(".pdf")
                try:
                    ocrmypdf.ocr(
                        str(pp), str(page_out),
                        language=language, deskew=True, clean=True,
                        optimize=1, force_ocr=True, progress_bar=False,
                    )
                    page_pdfs.append(page_out)
                except Exception as e:
                    print(f"OCR page failed: {e}")
                    # Fallback: image-only page
                    if REPORTLAB_AVAILABLE:
                        from reportlab.lib.pagesizes import A4
                        from reportlab.platypus import SimpleDocTemplate, Image as RLImg
                        from PIL import Image as PILImg
                        import io as _io
                        raw = pp.read_bytes()
                        img = PILImg.open(_io.BytesIO(raw))
                        W, H = A4
                        r = min(W/img.width, H/img.height)
                        doc_rl = SimpleDocTemplate(str(page_out), pagesize=A4,
                                                   leftMargin=0,rightMargin=0,
                                                   topMargin=0,bottomMargin=0)
                        doc_rl.build([RLImg(_io.BytesIO(raw), width=img.width*r, height=img.height*r)])
                        page_pdfs.append(page_out)

            # Merge using pikepdf (already installed as ocrmypdf dependency)
            if page_pdfs:
                merger = pikepdf.Pdf.new()
                for pp in page_pdfs:
                    try:
                        src = pikepdf.Pdf.open(str(pp))
                        merger.pages.extend(src.pages)
                    except Exception as e:
                        print(f"Merge page failed: {e}")
                merger.save(str(out_path))
            else:
                raise Exception("All pages failed OCR")

 except Exception as e:
        print(f"OCRmyPDF failed: {e}")
        # If it's already a single PDF, just return it instead of crashing Pillow
        if len(page_paths) == 1 and page_paths[0].suffix.lower() == ".pdf":
            return FileResponse(str(page_paths[0]), media_type="application/pdf", filename="fallback.pdf")
            
        # Final fallback: image-only PDF for JPG/PNG inputs
        if not REPORTLAB_AVAILABLE:
            raise HTTPException(500, f"PDF generation failed: {e}")
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Image as RLImg, PageBreak
        from PIL import Image as PILImg
        import io as _io
        story = []
        for i, pp in enumerate(page_paths):
            if pp.suffix.lower() == ".pdf": 
                continue # Skip PDFs to prevent Pillow UnidentifiedImageError
            raw = pp.read_bytes()
            img = PILImg.open(_io.BytesIO(raw))
            W, H = A4
            r = min(W/img.width, H/img.height)
            if i > 0: story.append(PageBreak())
            story.append(RLImg(_io.BytesIO(raw), width=img.width*r, height=img.height*r))
        doc_rl = SimpleDocTemplate(str(out_path), pagesize=A4,
                                   leftMargin=0,rightMargin=0,topMargin=0,bottomMargin=0)
        doc_rl.build(story)

    for pp in page_paths:
        background_tasks.add_task(rm, str(pp))

    return FileResponse(str(out_path), media_type="application/pdf", filename="searchable.pdf")


@app.post("/api/export/pdf")
async def export_pdf(background_tasks: BackgroundTasks,
                     text: str=Form(...), title: str=Form("Document"), author: str=Form("")):
    if not REPORTLAB_AVAILABLE: raise HTTPException(500, "reportlab not installed")
    sections = detect_structure(text)
    out = tmp(".pdf"); build_pdf(sections, title, author, str(out))
    background_tasks.add_task(rm, str(out))
    return FileResponse(str(out), media_type="application/pdf",
                        filename=f"{safe_name(title)}.pdf")


@app.post("/api/export/docx")
async def export_docx(background_tasks: BackgroundTasks,
                      text: str=Form(...), title: str=Form("Document"), author: str=Form("")):
    if not DOCX_AVAILABLE: raise HTTPException(500, "python-docx not installed")
    sections = detect_structure(text)
    out = tmp(".docx"); build_docx(sections, title, author, str(out))
    background_tasks.add_task(rm, str(out))
    return FileResponse(str(out),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=f"{safe_name(title)}.docx")


@app.post("/api/export/txt")
async def export_txt(background_tasks: BackgroundTasks,
                     text: str=Form(...), title: str=Form("Document")):
    out = tmp(".txt")
    header = f"{'='*70}\n  {title}\n  Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n{'='*70}\n\n"
    out.write_text(header + text, encoding="utf-8")
    background_tasks.add_task(rm, str(out))
    return FileResponse(str(out), media_type="text/plain", filename=f"{safe_name(title)}.txt")


@app.post("/api/ocrmypdf")
async def make_searchable(background_tasks: BackgroundTasks,
                          file: UploadFile=File(...), language: str=Form("eng")):
    if not OCRMYPDF_AVAILABLE: raise HTTPException(500, "ocrmypdf not installed")
    data = await file.read()
    ext  = Path(file.filename or "").suffix.lower() or ".jpg"

    inp = tmp(ext)
    inp.write_bytes(data)
    out = tmp("_searchable.pdf")

    try:
        # If input is an image, convert to PDF first via pdf2image / PIL
        if ext in IMAGE_EXTS or (file.content_type or "").startswith("image/"):
            if not PDF2IMAGE_AVAILABLE:
                raise HTTPException(400, "pdf2image not available for image input")
            from PIL import Image as PILImage
            img = PILImage.open(str(inp))
            if img.mode not in ("RGB", "L"): img = img.convert("RGB")
            img_pdf = tmp(".pdf")
            img.save(str(img_pdf), "PDF", resolution=150)
            del img
            import gc; gc.collect()
            ocrmypdf.ocr(str(img_pdf), str(out), language=language,
                         skip_text=False, deskew=False, clean=False,
                         optimize=0, output_type="pdf",
                         jobs=1)
            background_tasks.add_task(rm, str(img_pdf))
        else:
            ocrmypdf.ocr(str(inp), str(out), language=language,
                         skip_text=True, deskew=False, clean=False,
                         optimize=0, output_type="pdf",
                         jobs=1)

        background_tasks.add_task(rm, str(inp))
        background_tasks.add_task(rm, str(out))
        fname = Path(file.filename or "document").stem
        return FileResponse(str(out), media_type="application/pdf",
                            filename=f"{safe_name(fname)}_searchable.pdf")
    except Exception as e:
        background_tasks.add_task(rm, str(inp))
        raise HTTPException(500, f"OCRmyPDF failed: {str(e)[:200]}")
